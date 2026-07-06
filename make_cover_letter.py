from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches, Pt


OUT = Path(r"C:\Users\zannt\OneDrive\Papers sub\Nature\Nature climate change\Revision\cover_letter_draft.docx")


def add_para(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    return p


def add_highlighted_run(paragraph, text: str):
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return run


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(13)

    add_para(doc, "Cover letter draft", "Heading 1")
    p = add_para(doc)
    add_highlighted_run(p, "[Insert manuscript number if available]")
    add_para(doc, "6 July 2026")
    add_para(doc, "Dear Editor,")

    add_para(
        doc,
        "Please find enclosed our revised Brief Communication, \"Declining environmental predictability erodes reproduction in a tropical seabird\", for further consideration in Nature Climate Change.",
    )
    add_para(
        doc,
        "We are grateful for the opportunity to revise the manuscript. We have reformatted the paper as a Brief Communication and have addressed the editorial requests and referee comments in the revised files and in the accompanying response to reviewers.",
    )

    add_para(doc, "Response to editorial points", "Heading 2")
    bullets = [
        "Format: The manuscript has been revised as a Brief Communication with three display items.",
        "Title: The revised title is 80 characters, including spaces, and contains no punctuation.",
        "Abstract: The abstract is 65 words and summarizes the main findings without references.",
        "Main text: The main text is approximately 1,688 words.",
        "Headings and display items: Headings were checked for length, and all three display figures are cited in the main text. The display figures are supported by Extended Data items where needed.",
        "Extended Data and Supplementary Information: Extended Data Figs. 1-6 are captioned and cited, and Supplementary Tables S1-S9 are numbered consistently with manuscript and legend references.",
        "Data and code availability: A Code Availability statement is included. ",
        "End matter: Funding, acknowledgements, author contributions, and competing interests statements are included.",
    ]
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        if item.startswith("Data and code"):
            add_highlighted_run(
                p,
                "[Before submission, insert or verify a separate Data Availability statement and the final persistent repository DOI/details.]",
            )

    add_para(doc, "Submission metrics requested by the editors", "Heading 2")
    metrics = [
        "Text length: abstract, 65 words; main text, approximately 1,688 words.",
        "Methods length: approximately 1,089 words before Methods references; approximately 1,760 words including Methods references.",
        "Legend length: display-figure legends, approximately 590 words; Extended Data figure legends, approximately 630 words.",
        "References: 48 total, assuming main-text references 1-24 and Methods references 25-48. ",
        "Display items: 3 main figures and 0 main tables.",
        "Extended Data and supplementary items: 6 Extended Data figures and 9 Supplementary Tables.",
        "Estimated final size of current figure PDFs: 3.88 MB total (3.70 MiB) across seven one-page PDF files.",
    ]
    for item in metrics:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        if item.startswith("References"):
            add_highlighted_run(
                p,
                "[Verify that the main-reference bibliography is present in the final submission file; the current brief_ms.docx extraction shows the Methods references but not the main-reference entries.]",
            )

    add_para(
        doc,
        "We hope that the revised manuscript is now suitable for further consideration. Thank you for your time and for coordinating the review of our work.",
    )
    add_para(doc, "Sincerely,")
    add_para(doc, "Santiago Ortega, on behalf of all authors")
    p = add_para(doc)
    add_highlighted_run(p, "[Insert corresponding-author contact details, if desired]")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
