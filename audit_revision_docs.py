from __future__ import annotations

import re
from collections import defaultdict
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree
from pypdf import PdfReader


REVISION_DIR = Path(r"C:\Users\zannt\OneDrive\Papers sub\Nature\Nature climate change\Revision")

W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def read_docx_text(path: Path) -> str:
    parts: list[str] = []
    doc = Document(path)
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_docx_field_text(path: Path) -> str:
    """Read raw XML text so field/cross-reference display text is not missed."""
    texts: list[str] = []
    with ZipFile(path) as zf:
        for name in zf.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                if not (name == "word/document.xml" or name.startswith("word/header") or name.startswith("word/footer")):
                    continue
                root = etree.fromstring(zf.read(name))
                for node in root.xpath(".//w:t", namespaces=W_NS):
                    if node.text:
                        texts.append(node.text)
    return " ".join(texts)


def normalize_label(kind: str, n: str) -> str:
    return f"{kind} {int(n)}"


def extract_refs(text: str) -> dict[str, set[str]]:
    patterns = {
        "Figure": r"(?<!Extended Data )\b(?:Fig(?:ure)?\.?)\s+(\d+)",
        "Extended Data Fig.": r"\bExtended\s+Data\s+Fig(?:ure)?\.?\s+(\d+)",
        "Table": r"\bTable\s+(\d+)",
        "Supplementary Fig.": r"\bSupplementary\s+Fig(?:ure)?\.?\s*S?(\d+)",
        "Supplementary Table": r"\bSupplementary\s+Tables?\s*S?(\d+)(?:\s*[-–]\s*S?(\d+))?",
    }
    refs: dict[str, set[str]] = {}
    for label, pattern in patterns.items():
        values = set()
        for match in re.findall(pattern, text, flags=re.I):
            if isinstance(match, tuple):
                nums = [int(n) for n in match if n]
                if len(nums) == 2:
                    values.update(range(nums[0], nums[1] + 1))
                elif nums:
                    values.add(nums[0])
            else:
                values.add(int(match))
        refs[label] = {normalize_label(label, str(n)) for n in values}
    return refs


def extract_caption_lines(text: str) -> list[str]:
    lines = []
    for line in text.splitlines():
        clean = " ".join(line.split())
        if re.match(
            r"^(Fig(?:ure)?\.?\s+\d+|Extended Data Fig(?:ure)?\.?\s+\d+|Table\s+S?\d+|Supplementary (?:Fig(?:ure)?\.?|Table)\s+S?\d+)\b",
            clean,
            flags=re.I,
        ):
            lines.append(clean)
    return lines


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w'-]+\b", text))


def section_between(text: str, start_pat: str, end_pats: list[str]) -> str:
    m = re.search(start_pat, text, flags=re.I | re.M)
    if not m:
        return ""
    start = m.end()
    ends = [re.search(p, text[start:], flags=re.I | re.M) for p in end_pats]
    ends = [e.start() for e in ends if e]
    end = start + min(ends) if ends else len(text)
    return text[start:end].strip()


def main():
    docx_files = sorted(REVISION_DIR.glob("*.docx"))
    pdf_files = sorted(REVISION_DIR.glob("*.pdf"))

    all_refs: dict[str, dict[str, set[str]]] = {}
    all_caps: dict[str, list[str]] = {}
    texts: dict[str, str] = {}

    print("# Word files")
    for path in docx_files:
        text = read_docx_text(path)
        raw = read_docx_field_text(path)
        combined = text + "\n" + raw
        texts[path.name] = text
        all_refs[path.name] = extract_refs(combined)
        all_caps[path.name] = extract_caption_lines(text)
        print(f"\n## {path.name}")
        print(f"words_total: {word_count(text)}")
        print("captions:")
        for cap in all_caps[path.name]:
            print(f"  - {cap[:220]}")
        print("references:")
        for kind, refs in all_refs[path.name].items():
            if refs:
                print(f"  {kind}: {', '.join(sorted(refs, key=lambda x: int(re.search(r'\d+', x).group())))}")

    print("\n# Combined reference/caption sets")
    combined_refs = defaultdict(set)
    combined_caps = defaultdict(set)
    for refs in all_refs.values():
        for kind, values in refs.items():
            combined_refs[kind].update(values)
    for caps in all_caps.values():
        for cap in caps:
            if re.match(r"^Extended Data Fig", cap, flags=re.I):
                n = re.search(r"\d+", cap).group()
                combined_caps["Extended Data Fig."].add(normalize_label("Extended Data Fig.", n))
            elif re.match(r"^Supplementary Fig", cap, flags=re.I):
                n = re.search(r"\d+", cap).group()
                combined_caps["Supplementary Fig."].add(normalize_label("Supplementary Fig.", n))
            elif re.match(r"^Supplementary Table", cap, flags=re.I):
                n = re.search(r"\d+", cap).group()
                combined_caps["Supplementary Table"].add(normalize_label("Supplementary Table", n))
            elif re.match(r"^Table\s+S", cap, flags=re.I):
                n = re.search(r"\d+", cap).group()
                combined_caps["Supplementary Table"].add(normalize_label("Supplementary Table", n))
            elif re.match(r"^(?:Fig|Figure)", cap, flags=re.I):
                n = re.search(r"\d+", cap).group()
                combined_caps["Figure"].add(normalize_label("Figure", n))
            elif re.match(r"^Table", cap, flags=re.I):
                n = re.search(r"\d+", cap).group()
                combined_caps["Table"].add(normalize_label("Table", n))
    for kind in ["Figure", "Extended Data Fig.", "Table", "Supplementary Fig.", "Supplementary Table"]:
        refs = combined_refs[kind]
        caps = combined_caps[kind]
        if refs or caps:
            key = lambda x: int(re.search(r"\d+", x).group())
            print(f"{kind}: refs={', '.join(sorted(refs, key=key)) or '-'} | captions={', '.join(sorted(caps, key=key)) or '-'}")
            print(f"  referenced_without_caption={', '.join(sorted(refs - caps, key=key)) or '-'}")
            print(f"  caption_without_reference={', '.join(sorted(caps - refs, key=key)) or '-'}")

    print("\n# Brief Communication counts")
    brief = texts.get("brief_ms.docx", "")
    methods = texts.get("Methods.docx", "")
    figures = texts.get("Figures.docx", "")
    ed = texts.get("Extended Data Figs.docx", "")
    supp = texts.get("Supplementary material.docx", "")
    abstract = section_between(brief, r"^Abstract\s*$", [r"^\w", r"^Introduction", r"^Main"])
    refs_block = section_between(brief, r"^References\s*$", [r"^Methods references", r"^Acknowledg", r"^Author", r"^Competing"])
    print(f"brief_ms_total_words: {word_count(brief)}")
    print(f"methods_total_words: {word_count(methods)}")
    print(f"figures_legend_words: {word_count(figures)}")
    print(f"extended_data_legend_words: {word_count(ed)}")
    print(f"supplementary_material_words: {word_count(supp)}")
    ref_nums = sorted({int(n) for n in re.findall(r"^\s*(\d+)[\.\t ]", brief, flags=re.M)})
    print(f"numbered_references_in_brief_ms: {len(ref_nums)} ({ref_nums[0] if ref_nums else '-'}-{ref_nums[-1] if ref_nums else '-'})")

    print("\n# PDFs")
    total = 0
    for path in pdf_files:
        size = path.stat().st_size
        total += size
        pages = "?"
        try:
            pages = len(PdfReader(str(path)).pages)
        except Exception as exc:
            pages = f"? ({exc})"
        print(f"{path.name}: {size} bytes | {size / 1024 / 1024:.3f} MiB | pages={pages}")
    print(f"TOTAL_PDF_SIZE: {total} bytes | {total / 1024 / 1024:.3f} MiB | {total / 1000 / 1000:.3f} MB")


if __name__ == "__main__":
    main()
